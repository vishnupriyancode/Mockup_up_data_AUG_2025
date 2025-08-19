#!/usr/bin/env python3
"""
Project Workflow Documentation Generator
Creates a comprehensive .docx file with project workflow, screenshots, and diagrams
"""

import os
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
except ImportError:
    print("Error: python-docx library not found!")
    print("Please install it using: pip install python-docx")
    sys.exit(1)


class ProjectWorkflowDocumentation:
    """Generates comprehensive project workflow documentation in .docx format"""
    
    def __init__(self):
        self.document = Document()
        self.setup_document_styles()
        
    def setup_document_styles(self):
        """Setup document styles and formatting"""
        # Set document title
        title = self.document.add_heading('Mock Generation System - Project Workflow Documentation', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add subtitle
        subtitle = self.document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run(f'Generated on {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.italic = True
        
        self.document.add_page_break()
        
    def add_table_of_contents(self):
        """Add table of contents"""
        self.document.add_heading('Table of Contents', level=1)
        
        toc_items = [
            "1. Executive Summary",
            "2. Project Overview",
            "3. System Architecture",
            "4. Installation & Setup",
            "5. User Workflow",
            "6. Configuration Guide",
            "7. Usage Examples",
            "8. Output Examples",
            "9. Troubleshooting",
            "10. Advanced Features",
            "11. Project Benefits",
            "12. Conclusion"
        ]
        
        for item in toc_items:
            self.document.add_paragraph(item, style='List Bullet')
            
        self.document.add_page_break()
        
    def add_executive_summary(self):
        """Add executive summary section"""
        self.document.add_heading('1. Executive Summary', level=1)
        
        summary = self.document.add_paragraph()
        summary.add_run('The Mock Generation System is a powerful, enterprise-grade solution designed to generate high-quality mock JSON data for software testing, development, and data validation workflows. ').bold = False
        summary.add_run('This system addresses the critical need for realistic, consistent, and configurable test data across various development scenarios.').bold = False
        
        # Key Benefits
        self.document.add_heading('Key Benefits', level=2)
        benefits = [
            "🚀 Rapid mock data generation with probability-based scenarios",
            "🔄 Master template integration for consistent data structure",
            "📊 Multiple output formats (single, multiple, split files)",
            "🛡️ Comprehensive error handling and validation",
            "⚡ CLI interface for automation and CI/CD integration",
            "📁 Organized output management with timestamps"
        ]
        
        for benefit in benefits:
            self.document.add_paragraph(benefit, style='List Bullet')
            
    def add_project_overview(self):
        """Add project overview section"""
        self.document.add_heading('2. Project Overview', level=1)
        
        overview = self.document.add_paragraph()
        overview.add_run('The Mock Generation System was developed to solve common challenges in software development and testing:').bold = False
        
        challenges = [
            "Need for realistic test data that matches production scenarios",
            "Time-consuming manual data creation for different test cases",
            "Inconsistent data structures across different testing phases",
            "Lack of probability-based scenarios for comprehensive testing",
            "Difficulty in maintaining data consistency across team members"
        ]
        
        for challenge in challenges:
            self.document.add_paragraph(challenge, style='List Bullet')
            
        # Solution Overview
        self.document.add_heading('Our Solution', level=2)
        solution = self.document.add_paragraph()
        solution.add_run('A comprehensive system that provides:').bold = False
        
        solution_points = [
            "Automated mock data generation with configurable templates",
            "Probability-based scenarios (positive, negative, exclusion)",
            "Master template integration for consistent data structure",
            "Flexible output formats and file management",
            "Easy-to-use CLI interface for developers and testers"
        ]
        
        for point in solution_points:
            self.document.add_paragraph(point, style='List Bullet')
            
    def add_system_architecture(self):
        """Add system architecture section with diagrams"""
        self.document.add_heading('3. System Architecture', level=1)
        
        # Architecture Overview
        arch_overview = self.document.add_paragraph()
        arch_overview.add_run('The system follows a modular, layered architecture that ensures maintainability, extensibility, and robust operation.').bold = False
        
        # Architecture Diagram (Text-based)
        self.document.add_heading('System Architecture Diagram', level=2)
        
        # Create a table to represent the architecture
        arch_table = self.document.add_table(rows=4, cols=3)
        arch_table.style = 'Table Grid'
        
        # Headers
        arch_table.cell(0, 0).text = 'Layer'
        arch_table.cell(0, 1).text = 'Components'
        arch_table.cell(0, 2).text = 'Responsibilities'
        
        # User Interface Layer
        arch_table.cell(1, 0).text = 'User Interface'
        arch_table.cell(1, 1).text = '• CLI Interface\n• Configuration Files\n• Help System'
        arch_table.cell(1, 2).text = '• User interaction\n• Command processing\n• Input validation'
        
        # Core Processing Layer
        arch_table.cell(2, 0).text = 'Core Processing'
        arch_table.cell(2, 1).text = '• Data Generator\n• Probability Engine\n• Template Manager'
        arch_table.cell(2, 2).text = '• Data generation\n• Scenario processing\n• Template integration'
        
        # Data Management Layer
        arch_table.cell(3, 0).text = 'Data Management'
        arch_table.cell(3, 1).text = '• Input Parser\n• Output Formatter\n• File Manager'
        arch_table.cell(3, 2).text = '• Configuration loading\n• Output formatting\n• File operations'
        
        # Data Flow Diagram
        self.document.add_heading('Data Flow Process', level=2)
        
        flow_steps = [
            "1. User provides configuration (user_input.json)",
            "2. System loads master template (master.json)",
            "3. Configuration is validated and parsed",
            "4. Data generation engine processes models",
            "5. Probability scenarios are applied",
            "6. Output is formatted according to specifications",
            "7. Files are saved with timestamps and organization"
        ]
        
        for step in flow_steps:
            self.document.add_paragraph(step, style='List Number')
            
    def add_installation_setup(self):
        """Add installation and setup section"""
        self.document.add_heading('4. Installation & Setup', level=1)
        
        # Prerequisites
        self.document.add_heading('Prerequisites', level=2)
        prereq = self.document.add_paragraph()
        prereq.add_run('Before installing the Mock Generation System, ensure you have:').bold = False
        
        prereq_list = [
            "Python 3.6 or higher installed",
            "Git or download access to the project files",
            "Basic understanding of JSON configuration",
            "Command line interface access"
        ]
        
        for item in prereq_list:
            self.document.add_paragraph(item, style='List Bullet')
            
        # Installation Steps
        self.document.add_heading('Installation Steps', level=2)
        
        install_steps = [
            "Clone or download the project files to your local machine",
            "Navigate to the project directory: cd Mockup_up_data-main",
            "Install required dependencies: pip install python-docx",
            "Verify installation by running: python -m src.mockgen.cli --help",
            "Ensure user_input.json and master.json files are present"
        ]
        
        for step in install_steps:
            self.document.add_paragraph(step, style='List Number')
            
        # Quick Verification
        self.document.add_heading('Quick Verification', level=2)
        verify = self.document.add_paragraph()
        verify.add_run('Test that everything is working correctly:').bold = False
        
        verify_cmd = self.document.add_paragraph()
        verify_cmd.add_run('python -m src.mockgen.cli --enhanced').font.name = 'Courier New'
        verify_cmd.add_run(' - This should generate output files in the generated_outputs/ directory')
        
    def add_user_workflow(self):
        """Add user workflow section with step-by-step process"""
        self.document.add_heading('5. User Workflow', level=1)
        
        # Workflow Overview
        workflow_overview = self.document.add_paragraph()
        workflow_overview.add_run('The typical user workflow involves three main phases: Configuration, Generation, and Output Management.').bold = False
        
        # Phase 1: Configuration
        self.document.add_heading('Phase 1: Configuration', level=2)
        config_steps = [
            "Review and understand your data requirements",
            "Modify user_input.json with your specific data models",
            "Customize master.json template if needed",
            "Define probability scenarios (positive, negative, exclusion)",
            "Test configuration with small data sets"
        ]
        
        for step in config_steps:
            self.document.add_paragraph(step, style='List Number')
            
        # Phase 2: Generation
        self.document.add_heading('Phase 2: Generation', level=2)
        gen_steps = [
            "Choose the appropriate generation method",
            "Run CLI commands with desired parameters",
            "Monitor generation progress and output",
            "Verify generated data meets requirements",
            "Generate additional scenarios if needed"
        ]
        
        for step in gen_steps:
            self.document.add_paragraph(step, style='List Number')
            
        # Phase 3: Output Management
        self.document.add_heading('Phase 3: Output Management', level=2)
        output_steps = [
            "Review generated files in output directory",
            "Organize outputs by type and timestamp",
            "Integrate data into your testing workflow",
            "Archive or clean up old output files",
            "Document any custom configurations for team use"
        ]
        
        for step in output_steps:
            self.document.add_paragraph(step, style='List Number')
            
    def add_configuration_guide(self):
        """Add configuration guide section"""
        self.document.add_heading('6. Configuration Guide', level=1)
        
        # Configuration Files Overview
        config_overview = self.document.add_paragraph()
        config_overview.add_run('The system uses two main configuration files: user_input.json for your data models and master.json for the base template structure.').bold = False
        
        # User Input Configuration
        self.document.add_heading('user_input.json Configuration', level=2)
        user_config = self.document.add_paragraph()
        user_config.add_run('This file defines your data models and probability scenarios:').bold = False
        
        # Create configuration example table
        config_table = self.document.add_table(rows=5, cols=2)
        config_table.style = 'Table Grid'
        
        config_table.cell(0, 0).text = 'Model Type'
        config_table.cell(0, 1).text = 'Purpose'
        
        config_table.cell(1, 0).text = 'Model_1'
        config_table.cell(1, 1).text = 'Base model with standard data'
        
        config_table.cell(2, 0).text = 'Model_1_Positive'
        config_table.cell(2, 1).text = 'Valid data for positive testing'
        
        config_table.cell(3, 0).text = 'Model_1_Negative'
        config_table.cell(3, 1).text = 'Invalid data for error testing'
        
        config_table.cell(4, 0).text = 'Model_1_Exclusion'
        config_table.cell(4, 1).text = 'Data that should be filtered out'
        
        # Master Template Configuration
        self.document.add_heading('master.json Configuration', level=2)
        master_config = self.document.add_paragraph()
        master_config.add_run('The master template provides the base structure for all generated outputs:').bold = False
        
        master_points = [
            "Defines the overall data structure",
            "Provides default values for common fields",
            "Ensures consistency across all models",
            "Can be customized for your specific needs"
        ]
        
        for point in master_points:
            self.document.add_paragraph(point, style='List Bullet')
            
    def add_usage_examples(self):
        """Add usage examples section"""
        self.document.add_heading('7. Usage Examples', level=1)
        
        # Basic Usage
        self.document.add_heading('Basic Usage Examples', level=2)
        
        # Example 1
        self.document.add_heading('Example 1: Generate Enhanced Output for All Models', level=3)
        ex1_cmd = self.document.add_paragraph()
        ex1_cmd.add_run('python -m src.mockgen.cli --enhanced').font.name = 'Courier New'
        ex1_cmd.add_run(' - This generates output for all configured models using the master template')
        
        # Example 2
        self.document.add_heading('Example 2: Generate Output for Specific Model', level=3)
        ex2_cmd = self.document.add_paragraph()
        ex2_cmd.add_run('python -m src.mockgen.cli --enhanced --model Model_1').font.name = 'Courier New'
        ex2_cmd.add_run(' - This generates output only for Model_1')
        
        # Example 3
        self.document.add_heading('Example 3: Generate Multiple Records', level=3)
        ex3_cmd = self.document.add_paragraph()
        ex3_cmd.add_run('python -m src.mockgen.cli --model Model_1 --count 10').font.name = 'Courier New'
        ex3_cmd.add_run(' - This generates 10 records for Model_1')
        
        # Advanced Usage
        self.document.add_heading('Advanced Usage Examples', level=2)
        
        # Probability Generator
        self.document.add_heading('Probability Generator Usage', level=3)
        prob_examples = [
            "Generate all probability types: python -m src.mockgen.cli --probability --all --model Model_1 --wgs",
"Generate only positive scenarios: python -m src.mockgen.cli --probability --positive --model Model_1 --wgs",
"Generate with multiple records: python -m src.mockgen.cli --probability --all --model Model_1 --wgs --count 5",
"Generate multiple records: python -m src.mockgen.cli --probability --all --model Model_1 --wgs --count 3"
        ]
        
        for example in prob_examples:
            self.document.add_paragraph(example, style='List Bullet')
            
    def add_output_examples(self):
        """Add output examples section"""
        self.document.add_heading('8. Output Examples', level=1)
        
        # Output Structure
        self.document.add_heading('Output Structure Overview', level=2)
        output_overview = self.document.add_paragraph()
        output_overview.add_run('Generated outputs follow a consistent structure based on your configuration and the master template.').bold = False
        
        # Enhanced Output Example
        self.document.add_heading('Enhanced Output Example', level=2)
        enhanced_example = self.document.add_paragraph()
        enhanced_example.add_run('When using --enhanced mode, output includes the complete master template structure:').bold = False
        
        # Create output example table
        output_table = self.document.add_table(rows=6, cols=2)
        output_table.style = 'Table Grid'
        
        output_table.cell(0, 0).text = 'Field'
        output_table.cell(0, 1).text = 'Value'
        
        output_table.cell(1, 0).text = 'first_name'
        output_table.cell(1, 1).text = 'John'
        
        output_table.cell(2, 0).text = 'last_name'
        output_table.cell(2, 1).text = 'Smith'
        
        output_table.cell(3, 0).text = 'proc_cd'
        output_table.cell(3, 1).text = 'Vishnu'
        
        output_table.cell(4, 0).text = 'city'
        output_table.cell(4, 1).text = 'Hyderabad'
        
        output_table.cell(5, 0).text = 'country'
        output_table.cell(5, 1).text = 'United States'
        
        # File Naming Convention
        self.document.add_heading('File Naming Convention', level=2)
        naming_overview = self.document.add_paragraph()
        naming_overview.add_run('Generated files follow a consistent naming pattern:').bold = False
        
        naming_examples = [
            "enhanced_output_20241201_143022_123456Z.json",
            "Model_1_Positive_20241201_143022_123456Z.json",
            "Model_1_Negative_20241201_143022_123457Z.json",
            "Model_1_Exclusion_20241201_143022_123458Z.json"
        ]
        
        for example in naming_examples:
            ex_para = self.document.add_paragraph(example, style='List Bullet')
            ex_para.runs[0].font.name = 'Courier New'
            
    def add_troubleshooting(self):
        """Add troubleshooting section"""
        self.document.add_heading('9. Troubleshooting', level=1)
        
        # Common Issues
        self.document.add_heading('Common Issues and Solutions', level=2)
        
        # Issue 1
        self.document.add_heading('Issue 1: "Module not found" Error', level=3)
        issue1_desc = self.document.add_paragraph()
        issue1_desc.add_run('Problem: ').bold = True
        issue1_desc.add_run('Python cannot find the required modules')
        issue1_solution = self.document.add_paragraph()
        issue1_solution.add_run('Solution: ').bold = True
        issue1_solution.add_run('Ensure you are in the project root directory (Mockup_up_data-main) when running commands')
        
        # Issue 2
        self.document.add_heading('Issue 2: "File not found" Error', level=3)
        issue2_desc = self.document.add_paragraph()
        issue2_desc.add_run('Problem: ').bold = True
        issue2_desc.add_run('System cannot find configuration files')
        issue2_solution = self.document.add_paragraph()
        issue2_solution.add_run('Solution: ').bold = True
        issue2_solution.add_run('Verify that user_input.json and master.json exist in the current directory')
        
        # Issue 3
        self.document.add_heading('Issue 3: No Output Generated', level=3)
        issue3_desc = self.document.add_paragraph()
        issue3_desc.add_run('Problem: ').bold = True
        issue3_desc.add_run('Commands run successfully but no output files are created')
        issue3_solution = self.document.add_paragraph()
        issue3_solution.add_run('Solution: ').bold = True
        issue3_solution.add_run('Check that the generated_outputs/ folder exists and is writable')
        
        # Debug Commands
        self.document.add_heading('Debug Commands', level=2)
        debug_commands = [
            "python -m src.mockgen.cli --help - Shows all available options",
            "python -m src.mockgen.cli --list - Lists available models",
            "python -m src.mockgen.cli --init - Creates template configuration files"
        ]
        
        for command in debug_commands:
            self.document.add_paragraph(command, style='List Bullet')
            
    def add_advanced_features(self):
        """Add advanced features section"""
        self.document.add_heading('10. Advanced Features', level=1)
        
        # Split File Generation
        self.document.add_heading('Split File Generation', level=2)
        split_desc = self.document.add_paragraph()
        split_desc.add_run('Generate separate files for each record or model:').bold = False
        
        split_example = self.document.add_paragraph()
        split_example.add_run('python -m src.mockgen.cli --enhanced --output-format split --count 5').font.name = 'Courier New'
        split_example.add_run(' - Creates 5 separate files, one for each record')
        
        # Master Template Integration
        self.document.add_heading('Master Template Integration', level=2)
        template_desc = self.document.add_paragraph()
        template_desc.add_run('The system automatically merges your configuration with the master template:').bold = False
        
        template_benefits = [
            "Consistent data structure across all outputs",
            "Automatic field population from template",
            "Flexible override of template values",
            "Professional, standardized output format"
        ]
        
        for benefit in template_benefits:
            self.document.add_paragraph(benefit, style='List Bullet')
            
        # Probability Scenarios
        self.document.add_heading('Probability Scenarios', level=2)
        prob_desc = self.document.add_paragraph()
        prob_desc.add_run('Generate different types of test scenarios:').bold = False
        
        prob_types = [
            "Positive: Valid data for normal testing",
            "Negative: Invalid data for error handling testing",
            "Exclusion: Data that should be filtered out"
        ]
        
        for prob_type in prob_types:
            self.document.add_paragraph(prob_type, style='List Bullet')
            
    def add_project_benefits(self):
        """Add project benefits section"""
        self.document.add_heading('11. Project Benefits', level=1)
        
        # Developer Benefits
        self.document.add_heading('Benefits for Developers', level=2)
        dev_benefits = [
            "🚀 Time savings through automated data generation",
            "🔄 Consistent test data across development phases",
            "🛡️ Better testing coverage with probability scenarios",
            "⚡ Faster development and testing cycles",
            "📊 Realistic data that matches production scenarios"
        ]
        
        for benefit in dev_benefits:
            self.document.add_paragraph(benefit, style='List Bullet')
            
        # Team Benefits
        self.document.add_heading('Benefits for Teams', level=2)
        team_benefits = [
            "👥 Standardized data generation across team members",
            "🤝 Easy sharing of data configurations",
            "📋 Consistent testing standards and procedures",
            "🔄 Simplified onboarding for new team members",
            "📊 Unified approach to test data management"
        ]
        
        for benefit in team_benefits:
            self.document.add_paragraph(benefit, style='List Bullet')
            
        # Organization Benefits
        self.document.add_heading('Benefits for Organizations', level=2)
        org_benefits = [
            "💰 Cost reduction through faster development cycles",
            "🎯 Improved software quality and reliability",
            "🛡️ Reduced risk through comprehensive testing",
            "📈 Scalable testing capabilities as projects grow",
            "🏆 Professional, enterprise-grade testing tools"
        ]
        
        for benefit in org_benefits:
            self.document.add_paragraph(benefit, style='List Bullet')
            
    def add_conclusion(self):
        """Add conclusion section"""
        self.document.add_heading('12. Conclusion', level=1)
        
        conclusion = self.document.add_paragraph()
        conclusion.add_run('The Mock Generation System provides a comprehensive solution for generating high-quality mock data for software development and testing. ').bold = False
        conclusion.add_run('With its advanced features, flexible configuration, and professional output formats, it significantly improves the efficiency and quality of development workflows.').bold = False
        
        # Key Takeaways
        self.document.add_heading('Key Takeaways', level=2)
        takeaways = [
            "The system is designed to be easy to use while providing powerful capabilities",
            "Master template integration ensures consistent, professional outputs",
            "Probability-based scenarios enable comprehensive testing coverage",
            "Flexible output formats support various development and testing needs",
            "The CLI interface enables automation and integration with existing workflows"
        ]
        
        for takeaway in takeaways:
            self.document.add_paragraph(takeaway, style='List Bullet')
            
        # Getting Started
        self.document.add_heading('Getting Started', level=2)
        getting_started = self.document.add_paragraph()
        getting_started.add_run('To begin using the Mock Generation System:').bold = False
        
        start_steps = [
            "1. Review the configuration files and understand the structure",
            "2. Start with simple models and gradually expand complexity",
            "3. Experiment with different output formats and scenarios",
            "4. Integrate the system into your existing development workflow",
            "5. Customize configurations to match your specific needs"
        ]
        
        for step in start_steps:
            self.document.add_paragraph(step, style='List Number')
            
        # Final Note
        final_note = self.document.add_paragraph()
        final_note.add_run('The system is designed to grow with your needs. Start simple and expand as you become more familiar with its capabilities. ').bold = False
        final_note.add_run('For additional support, refer to the built-in help system and comprehensive documentation.').bold = False
        
    def add_appendix(self):
        """Add appendix with additional information"""
        self.document.add_page_break()
        self.document.add_heading('Appendix A: Command Reference', level=1)
        
        # CLI Commands Reference
        self.document.add_heading('Enhanced System CLI Commands', level=2)
        
        cli_commands = [
            ("--enhanced", "Use enhanced mode with master template integration"),
            ("--model MODEL", "Generate output for specific model"),
            ("--models MODEL1 MODEL2", "Generate output for multiple specific models"),
            ("--count N", "Generate N number of records"),
            ("--output-format FORMAT", "Set output format (single/multiple/split)"),
            ("--split", "Write each output to separate file"),
            ("--config FILE", "Path to custom configuration file"),
            ("--master FILE", "Path to custom master template file"),
            ("--init", "Create template configuration files"),
            ("--legacy", "Force legacy mode for backward compatibility")
        ]
        
        cli_table = self.document.add_table(rows=len(cli_commands) + 1, cols=2)
        cli_table.style = 'Table Grid'
        
        cli_table.cell(0, 0).text = 'Command Option'
        cli_table.cell(0, 1).text = 'Description'
        
        for i, (command, description) in enumerate(cli_commands, 1):
            cli_table.cell(i, 0).text = command
            cli_table.cell(i, 1).text = description
            
        # Probability Generator Commands
        self.document.add_heading('Probability Generator Commands', level=2)
        
        prob_commands = [
            ("--all", "Generate all probability types (positive, negative, exclusion)"),
            ("--positive", "Generate only positive probability outputs"),
            ("--negative", "Generate only negative probability outputs"),
            ("--exclusion", "Generate only exclusion probability outputs"),
            ("--model MODEL", "Generate outputs for specific model only"),
            ("--count COUNT", "Number of records to generate per probability type"),
            ("--split", "Generate separate files for each record"),
            ("--config FILE", "Path to custom configuration file"),
            ("--output-dir DIR", "Custom output directory for generated files"),
            ("--list", "List available models and their probability types")
        ]
        
        prob_table = self.document.add_table(rows=len(prob_commands) + 1, cols=2)
        prob_table.style = 'Table Grid'
        
        prob_table.cell(0, 0).text = 'Command Option'
        prob_table.cell(0, 1).text = 'Description'
        
        for i, (command, description) in enumerate(prob_commands, 1):
            prob_table.cell(i, 0).text = command
            prob_table.cell(i, 1).text = description
            
    def generate_document(self):
        """Generate the complete document"""
        print("Generating comprehensive project workflow documentation...")
        
        # Add all sections
        self.add_table_of_contents()
        self.add_executive_summary()
        self.add_project_overview()
        self.add_system_architecture()
        self.add_installation_setup()
        self.add_user_workflow()
        self.add_configuration_guide()
        self.add_usage_examples()
        self.add_output_examples()
        self.add_troubleshooting()
        self.add_advanced_features()
        self.add_project_benefits()
        self.add_conclusion()
        self.add_appendix()
        
        # Create output directory
        output_dir = Path("docs/project-workflow")
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # Save document
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_file = output_dir / f"Mock_Generation_System_Workflow_Documentation_{timestamp}.docx"
        
        self.document.save(str(output_file))
        print(f"Document generated successfully: {output_file}")
        
        return output_file


def main():
    """Main function to generate the documentation"""
    try:
        # Create documentation generator
        doc_generator = ProjectWorkflowDocumentation()
        
        # Generate the document
        output_file = doc_generator.generate_document()
        
        print("\n" + "="*60)
        print("🎉 PROJECT WORKFLOW DOCUMENTATION GENERATED SUCCESSFULLY!")
        print("="*60)
        print(f"📄 Document saved to: {output_file}")
        print(f"📁 Location: docs/project-workflow/")
        print(f"💡 Open the .docx file with Microsoft Word or compatible software")
        print("="*60)
        
    except Exception as e:
        print(f"❌ Error generating documentation: {e}")
        print("Please ensure you have the required dependencies installed:")
        print("pip install python-docx")
        sys.exit(1)


if __name__ == "__main__":
    main()
